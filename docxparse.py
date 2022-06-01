#!/usr/bin/env python

import sys
import docx
import json
import copy
from docx2python import docx2python as dx2py

question = {
    "name": "",
    "answers": [],
    "answerIndex": 0
}

questions = { "questions": [] }

reading_name = False


def ns_tag_name(node, name):
    if node.nsmap and node.prefix:
        return "{{{:s}}}{:s}".format(node.nsmap[node.prefix], name)
    return name


def descendants(node, desc_strs):
    if node is None:
        return []
    if not desc_strs:
        return [node]
    ret = {}
    for child_str in desc_strs[0]:
        for child in node.iterchildren(ns_tag_name(node, child_str)):
            descs = descendants(child, desc_strs[1:])
            if not descs:
                continue
            cd = ret.setdefault(child_str, [])
            if isinstance(descs, list):
                cd.extend(descs)
            else:
                cd.append(descs)
    return ret


def simplified_descendants(desc_dict):
    ret = []
    for vs in desc_dict.values():
        for v in vs:
            if isinstance(v, dict):
                ret.extend(simplified_descendants(v))
            else:
                ret.append(v)
    return ret


def process_list_data(attrs, dx2py_elem):
    desc = simplified_descendants(attrs)[0]
    level = int(desc.attrib[ns_tag_name(desc, "val")])
    elem = [i for i in dx2py_elem[0].split("\t") if i][0]#.rstrip(")")
    return "    " * level + elem + " "


def main(*argv):
    fname = r"./fiiso_pagina_intento_2.docx"
    docd = docx.Document(fname)
    docdpy = dx2py(fname)
    dr = docdpy.docx_reader
    docdpy_runs = docdpy.document_runs[0][0][0]
    if len(docd.paragraphs) != len(docdpy_runs):
        print("Lengths don't match. Abort")
        return -1
    subnode_tags = (("pPr",), ("numPr",), ("ilvl",))  # (("pPr",), ("numPr",), ("ilvl", "numId"))  # numId is for matching elements from word/numbering.xml
    highlighted_tags = (("r",), ("rPr",), ("highlight",))
    for idx, (par, l) in enumerate(zip(docd.paragraphs, docdpy_runs)):
        numbered_attrs = descendants(par._element, subnode_tags)
        highlighted_attrs = descendants(par._element, highlighted_tags)
        if numbered_attrs:
            question["answers"].append({"name": par.text})

            if highlighted_attrs:
                question["answerIndex"] = len(question["answers"]) - 1
                print("Correct answer")
        else:
            # Check if name is not empty
            if par.text:
                # Check if it starts with a letter and a dot, if so it's answer
                if par.text[0].isalpha() and par.text[1] == ".":
                    print("Answer:", par.text)
                    question["answers"].append({"name": par.text})

                    if highlighted_attrs:
                        question["answerIndex"] = len(question["answers"]) - 1
                        print("Correct answer")
                else:
                    print("Question name: " + par.text)

                    if(len(question["answers"]) > 0):
                        questions["questions"].append(copy.deepcopy(question))
                        question["answers"].clear()
                        question["answerIndex"] = 0
                        question["name"] = par.text
                    else:
                        print("Question with no answers not included!")
                        question["name"] += "<br>" + par.text

                    


if __name__ == "__main__":
    print("Python {:s} {:03d}bit on {:s}\n".format(" ".join(elem.strip() for elem in sys.version.split("\n")),
                                                   64 if sys.maxsize > 0x100000000 else 32, sys.platform))
    rc = main(*sys.argv[1:])
    print("\nDone. Parsed " + str(len(questions["questions"])) + " questions.")

    # save json file
    with open("out.json", "w", encoding="utf-8") as f:
        json.dump(questions, f, indent=4, ensure_ascii=False)

    sys.exit(rc)